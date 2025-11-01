from docx import Document
import re
import json
from typing import List, Dict, Any, Optional
import os
from pydantic import BaseModel, Field
from langchain_google_genai import ChatGoogleGenerativeAI
import uuid
from datetime import datetime
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware

# Setup API key
os.environ["GOOGLE_API_KEY"] = "AIzaSyACY8UxSWJCrNN4vWmdiRC5swOxPTmkmzM"

# Initialize LLM
llm = ChatGoogleGenerativeAI(model="gemini-2.5-flash")

### Meta data generation area

# Regex patterns for placeholders
placeholder_patterns = [
    r"\[[^\]]+\]",                           # [Client Name]
    r"\[\[[^\]]+\]\]",                       # [[Client Name]]
    r"<[^>]+>",                               # <Client Name>
    r"\{[^}]+\}",                            # {Client Name}
    r"_{3,}",                                # _____ (3 or more underscores)
    r"\b[A-Z]+(?:_[A-Z]+)+\b",              # CLIENT_NAME or DATE_OF_SAFE
    r"\$\{[^}]+\}",                          # ${AMOUNT}
    r"%[A-Za-z_]+%",                         # %DATE% or %CLIENT_NAME%
    r"\$[A-Z_]+\$",                          # $CLIENT_NAME$
    r"\[[A-Z]{2,}\]"                         # [FN], [LN], [DOE]
]

# Compile patterns into one combined regex
combined_pattern = re.compile('|'.join(placeholder_patterns))

# Sentence splitting regex
sentence_endings = re.compile(r'[.!?]+[\s\n]+|[\n]{2,}')

def extract_sentences(text):
    """Extract sentences from text, preserving their positions"""
    sentences = []
    start = 0
    
    for match in sentence_endings.finditer(text):
        end = match.end()
        sentence = text[start:end].strip()
        if sentence:
            sentences.append({
                'text': sentence,
                'start': start,
                'end': end
            })
        start = end
    
    # Add last sentence if text doesn't end with sentence ending
    if start < len(text):
        sentence = text[start:].strip()
        if sentence:
            sentences.append({
                'text': sentence,
                'start': start,
                'end': len(text)
            })
    
    return sentences

def find_sentence_context(sentences, match_pos, match_length):
    """Find the sentence containing the match and the sentences before/after"""
    match_end = match_pos + match_length
    sentence_before = None
    sentence_with_match = None
    sentence_after = None
    
    for i, sent in enumerate(sentences):
        if sent['start'] <= match_pos < sent['end']:
            sentence_with_match = sent['text']
            if i > 0:
                sentence_before = sentences[i-1]['text']
            if i < len(sentences) - 1:
                sentence_after = sentences[i+1]['text']
            break
    
    return sentence_before, sentence_with_match, sentence_after

def get_paragraph_text_with_context(doc, para_idx):
    """Get paragraph text with surrounding context"""
    context_before = ""
    context_after = ""
    
    # Get previous paragraph
    if para_idx > 0:
        prev_text = ''.join([run.text for run in doc.paragraphs[para_idx - 1].runs])
        context_before = prev_text.strip()
    
    # Get next paragraph
    if para_idx < len(doc.paragraphs) - 1:
        next_text = ''.join([run.text for run in doc.paragraphs[para_idx + 1].runs])
        context_after = next_text.strip()
    
    return context_before, context_after

def estimate_page_number(paragraph_index, total_paragraphs, avg_paragraphs_per_page=20):
    """
    Estimate page number based on paragraph position.
    Note: This is an approximation. Actual page numbers depend on:
    - Font sizes, margins, spacing
    - Images, tables, headers, footers
    - Word's rendering engine
    
    For accurate page numbers, you'd need to:
    1. Use COM automation (win32com on Windows)
    2. Convert to PDF and extract page info
    3. Use Aspose.Words or similar library
    """
    return min(int(paragraph_index / avg_paragraphs_per_page) + 1, max(1, total_paragraphs // avg_paragraphs_per_page))

def collect_placeholder_metadata(doc_path):
    """Collect comprehensive metadata for all placeholders in document"""
    doc = Document(doc_path)
    all_metadata = []
    placeholder_counter = 0  # Counter for unique IDs
    
    # Calculate total paragraphs for page estimation
    total_paragraphs = len([p for p in doc.paragraphs if ''.join([r.text for r in p.runs]).strip()])
    
    # Process paragraphs
    for para_idx, paragraph in enumerate(doc.paragraphs):
        full_text = ''.join([run.text for run in paragraph.runs])
        
        if not full_text.strip():
            continue
        
        sentences = extract_sentences(full_text)
        
        # Find all matches with their positions
        for match in combined_pattern.finditer(full_text):
            match_text = match.group()
            match_start = match.start()
            match_end = match.end()
            
            # Get sentence context
            sentence_before, sentence_with_match, sentence_after = find_sentence_context(
                sentences, match_start, len(match_text)
            )
            
            # Get paragraph context
            para_context_before, para_context_after = get_paragraph_text_with_context(doc, para_idx)
            
            # Extract surrounding text (100 chars before and after)
            context_start = max(0, match_start - 100)
            context_end = min(len(full_text), match_end + 100)
            surrounding_text = full_text[context_start:context_end]
            match_position_in_context = match_start - context_start
            
            # Get run information
            char_pos = 0
            run_info = []
            for run_idx, run in enumerate(paragraph.runs):
                run_start = char_pos
                run_end = char_pos + len(run.text)
                
                if run_start <= match_start < run_end or run_start < match_end <= run_end:
                    run_info.append({
                        'run_index': run_idx,
                        'text': run.text,
                        'bold': run.bold,
                        'italic': run.italic,
                        'underline': run.underline,
                        'font_name': run.font.name if run.font and run.font.name else None,
                        'font_size': str(run.font.size) if run.font and run.font.size else None,
                    })
                
                char_pos = run_end
            
            # Estimate page number
            estimated_page = estimate_page_number(para_idx, total_paragraphs)
            
            # Generate unique ID
            placeholder_counter += 1
            unique_id = f"PLACEHOLDER_{placeholder_counter:04d}"
            
            metadata = {
                'unique_id': unique_id,
                'match': match_text,
                'match_type': 'paragraph',
                'paragraph_index': para_idx,
                'estimated_page_number': estimated_page,
                'note_about_page': 'Page number is estimated. Actual pages depend on Word rendering.',
                'position_in_paragraph': match_start,
                'position_in_document': None,  # Can be calculated if needed
                'sentence_before': sentence_before,
                'sentence_with_match': sentence_with_match,
                'sentence_after': sentence_after,
                'paragraph_context_before': para_context_before[:200] if para_context_before else None,
                'paragraph_context_after': para_context_after[:200] if para_context_after else None,
                'surrounding_text': surrounding_text,
                'match_position_in_context': match_position_in_context,
                'run_information': run_info,
                'paragraph_alignment': str(paragraph.alignment) if paragraph.alignment else None,
                'paragraph_style': paragraph.style.name if paragraph.style else None,
                'full_paragraph_text': full_text[:500],  # First 500 chars
                'paragraph_length': len(full_text),
                # LLM and filling fields
                'llm_context': None,  # Will be populated by LLM with context about what to fill
                'is_filled': False,  # Whether this placeholder has been filled
                'value': None,  # The value that will replace the placeholder
            }
            
            all_metadata.append(metadata)
    
    # Process tables
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para_idx_in_cell, paragraph in enumerate(cell.paragraphs):
                    full_text = ''.join([run.text for run in paragraph.runs])
                    
                    if not full_text.strip():
                        continue
                    
                    sentences = extract_sentences(full_text)
                    
                    for match in combined_pattern.finditer(full_text):
                        match_text = match.group()
                        match_start = match.start()
                        match_end = match.end()
                        
                        sentence_before, sentence_with_match, sentence_after = find_sentence_context(
                            sentences, match_start, len(match_text)
                        )
                        
                        context_start = max(0, match_start - 100)
                        context_end = min(len(full_text), match_end + 100)
                        surrounding_text = full_text[context_start:context_end]
                        match_position_in_context = match_start - context_start
                        
                        run_info = []
                        char_pos = 0
                        for run_idx, run in enumerate(paragraph.runs):
                            run_start = char_pos
                            run_end = char_pos + len(run.text)
                            
                            if run_start <= match_start < run_end or run_start < match_end <= run_end:
                                run_info.append({
                                    'run_index': run_idx,
                                    'text': run.text,
                                    'bold': run.bold,
                                    'italic': run.italic,
                                    'underline': run.underline,
                                    'font_name': run.font.name if run.font and run.font.name else None,
                                    'font_size': str(run.font.size) if run.font and run.font.size else None,
                                })
                            
                            char_pos = run_end
                        
                        # Generate unique ID for table placeholders too
                        placeholder_counter += 1
                        unique_id = f"PLACEHOLDER_{placeholder_counter:04d}"
                        
                        metadata = {
                            'unique_id': unique_id,
                            'match': match_text,
                            'match_type': 'table',
                            'table_index': table_idx,
                            'row_index': row_idx,
                            'cell_index': cell_idx,
                            'paragraph_index_in_cell': para_idx_in_cell,
                            'position_in_paragraph': match_start,
                            'sentence_before': sentence_before,
                            'sentence_with_match': sentence_with_match,
                            'sentence_after': sentence_after,
                            'surrounding_text': surrounding_text,
                            'match_position_in_context': match_position_in_context,
                            'run_information': run_info,
                            'paragraph_style': paragraph.style.name if paragraph.style else None,
                            'full_paragraph_text': full_text[:500],
                            # LLM and filling fields
                            'llm_context': None,  # Will be populated by LLM with context about what to fill
                            'is_filled': False,  # Whether this placeholder has been filled
                            'value': None,  # The value that will replace the placeholder
                        }
                        
                        all_metadata.append(metadata)
    
    return all_metadata


def generate_placeholder_metadata(doc_path: str, output_file: Optional[str] = None, verbose: bool = False) -> Dict[str, Any]:
    """
    Generate placeholder metadata for a Word document.
    
    Args:
        doc_path: Path to the .docx file
        output_file: Optional path to save JSON output. If None, doesn't save to file.
        verbose: If True, prints summary and detailed information to console.
    
    Returns:
        Dictionary containing:
            - 'summary': Summary statistics and document info
            - 'placeholders': List of all placeholder metadata dictionaries
    """
    # Collect metadata
    metadata = collect_placeholder_metadata(doc_path)
    
    # Load document for statistics
    doc = Document(doc_path)
    total_paragraphs = len([p for p in doc.paragraphs if ''.join([r.text for r in p.runs]).strip()])
    
    # Create summary report
    summary = {
        'document_path': doc_path,
        'total_placeholders_found': len(metadata),
        'unique_placeholders': sorted(set(m['match'] for m in metadata)),
        'unique_placeholder_count': len(set(m['match'] for m in metadata)),
        'placeholders_by_type': {},
        'placeholders_by_paragraph': {},
        'document_statistics': {
            'total_paragraphs': total_paragraphs,
            'total_tables': len(doc.tables),
            'estimated_total_pages': estimate_page_number(total_paragraphs, total_paragraphs)
        }
    }
    
    # Count by match type
    for meta in metadata:
        match = meta['match']
        if match not in summary['placeholders_by_type']:
            summary['placeholders_by_type'][match] = 0
        summary['placeholders_by_type'][match] += 1
    
    # Group by paragraph
    for meta in metadata:
        if meta['match_type'] == 'paragraph':
            para_idx = meta['paragraph_index']
            if para_idx not in summary['placeholders_by_paragraph']:
                summary['placeholders_by_paragraph'][para_idx] = []
            summary['placeholders_by_paragraph'][para_idx].append(meta['match'])
    
    # Create output data
    output_data = {
        'summary': summary,
        'placeholders': metadata
    }
    
    # Print summary if verbose
    if verbose:
        print("="*80)
        print("PLACEHOLDER METADATA SUMMARY")
        print("="*80)
        print(f"\nDocument: {doc_path}")
        print(f"Total placeholders found: {len(metadata)}")
        print(f"Unique placeholders: {summary['unique_placeholder_count']}")
        print(f"\nUnique matches:")
        for match in summary['unique_placeholders']:
            count = summary['placeholders_by_type'][match]
            print(f"  - {match}: {count} occurrence(s)")
        
        print(f"\nDocument Statistics:")
        print(f"  - Total paragraphs: {summary['document_statistics']['total_paragraphs']}")
        print(f"  - Total tables: {summary['document_statistics']['total_tables']}")
        print(f"  - Estimated pages: {summary['document_statistics']['estimated_total_pages']} (approximate)")
        
        # Print detailed metadata for first few matches
        print("\n" + "="*80)
        print("DETAILED METADATA FOR FIRST 3 MATCHES:")
        print("="*80)
        for i, meta in enumerate(metadata[:3]):
            print(f"\n--- Match {i+1}: {meta['match']} ---")
            print(f"  Unique ID: {meta['unique_id']}")
            print(f"  Type: {meta['match_type']}")
            if 'estimated_page_number' in meta:
                print(f"  Estimated Page: {meta['estimated_page_number']}")
            print(f"  Paragraph Index: {meta.get('paragraph_index', 'N/A')}")
            print(f"  Is Filled: {meta['is_filled']}")
            print(f"  Value: {meta['value'] if meta['value'] else 'Not yet filled'}")
            print(f"  Sentence with match: {meta['sentence_with_match'][:100] if meta['sentence_with_match'] else 'N/A'}...")
            if meta['sentence_before']:
                print(f"  Sentence before: {meta['sentence_before'][:80]}...")
            if meta['sentence_after']:
                print(f"  Sentence after: {meta['sentence_after'][:80]}...")
            print(f"  Surrounding text: {meta['surrounding_text'][:100]}...")
    
    # Save to file if output_file is provided
    if output_file:
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(output_data, f, indent=2, default=str, ensure_ascii=False)
        if verbose:
            print(f"\n\nFull metadata saved to: {output_file}")
            print(f"  - Summary statistics included")
            print(f"  - Detailed metadata for all {len(metadata)} placeholders")
    
    return output_data



### ************ LLM CONTEXT AUGUMENTATION AREA ************



class PlaceholderContext(BaseModel):
    """Detailed context about a placeholder in a legal document."""
    
    placeholder_id: str = Field(description="Unique ID of the placeholder")
    llm_context: str = Field(
        description="Comprehensive context including: purpose/use in document, questions to ask user, "
                   "what it represents, possible value types, typical ranges/constraints, examples, "
                   "and any other relevant information to help fill this placeholder accurately"
    )


class PlaceholderContextsList(BaseModel):
    """List of placeholder contexts to fill."""
    
    contexts: list[PlaceholderContext] = Field(
        description="List of all placeholder contexts to fill"
    )


def generate_placeholder_contexts(metadata_json_path: str, docx_path: str) -> list[dict]:
    """
    Generate comprehensive LLM context for each placeholder in the document.
    
    Args:
        metadata_json_path: Path to the placeholder_metadata.json file
        docx_path: Path to the original .docx file
    
    Returns:
        List of dictionaries with 'placeholder_id' and 'llm_context' fields
    """
    # Load the metadata JSON
    with open(metadata_json_path, 'r', encoding='utf-8') as f:
        metadata_data = json.load(f)
    
    # Extract all placeholders
    placeholders = metadata_data['placeholders']
    
    # Load the document to get full text context
    doc = Document(docx_path)
    
    # Extract full document text for context
    full_document_text = []
    for para in doc.paragraphs:
        text = ''.join([run.text for run in para.runs])
        if text.strip():
            full_document_text.append(text)
    
    document_text_sample = '\n\n'.join(full_document_text[:])  # First 50 paragraphs for context
    
    # Prepare the detailed prompt with all metadata for the LLM
    # Limit context size to avoid token limits
    placeholder_json = json.dumps(placeholders, indent=2)
    
    prompt = f"""You are an expert legal document analyst. Analyze the following placeholders from a legal document and provide comprehensive context for each one.

    DOCUMENT CONTEXT (first 50 paragraphs):
    {document_text_sample}

    PLACEHOLDER METADATA (in JSON format):
    {placeholder_json}

    For EACH placeholder, provide detailed context including:
    1. **Purpose in Document**: What role does this placeholder serve in the document?
    2. **Use Case**: When and why would someone need to fill this?
    3. **Questions to Ask User**: What specific questions should be asked to get the correct value?
    4. **Value Description**: What does this placeholder represent?
    5. **Value Types**: What types of data are expected? (text, number, date, currency, name, etc.)
    6. **Constraints/Ranges**: Any length limits, format requirements, or valid ranges?
    7. **Examples**: Provide 2-3 realistic example values
    8. **Legal/Business Context**: Any legal or business implications of this field?
    9. **Related Fields**: Are there other placeholders this relates to?
    10. **Validation Rules**: What validation should be applied?

    Be specific and practical. Consider this is a legal SAFE (Simple Agreement for Future Equity) document.

    Return the results as a structured JSON array where each element has:
    - placeholder_id: The unique_id from the metadata
    - llm_context: A comprehensive, detailed paragraph covering all the above aspects

    Output format should be valid JSON only.
    """
    
    # Use structured output to get the response
    structured_llm = llm.with_structured_output(PlaceholderContextsList)
    
    try:
        response = structured_llm.invoke([{"role": "user", "content": prompt}])
        
        # Convert to list of dicts
        result = [
            {
                "placeholder_id": context.placeholder_id,
                "llm_context": context.llm_context
            }
            for context in response.contexts
        ]
        
        return result
        
    except Exception as e:
        print(f"Error generating contexts: {e}")
        # Fallback: try without structured output
        response = llm.invoke([{"role": "user", "content": prompt}])
        print("Raw response:")
        print(response.content)
        return []


def update_metadata_with_contexts(metadata_json_path: str, contexts: list[dict], output_path: str = None):
    """
    Update the metadata JSON with LLM contexts.
    
    Args:
        metadata_json_path: Path to input metadata JSON
        contexts: List of dicts with 'placeholder_id' and 'llm_context'
        output_path: Optional output path (defaults to overwriting input file)
    """
    # Load the existing metadata
    with open(metadata_json_path, 'r', encoding='utf-8') as f:
        metadata_data = json.load(f)
    
    # Create a lookup dictionary for contexts by placeholder_id
    context_lookup = {ctx['placeholder_id']: ctx['llm_context'] for ctx in contexts}
    
    # Update each placeholder's llm_context field
    updated_count = 0
    for placeholder in metadata_data['placeholders']:
        if placeholder['unique_id'] in context_lookup:
            placeholder['llm_context'] = context_lookup[placeholder['unique_id']]
            updated_count += 1
    
    # Save the updated metadata
    output_file = output_path if output_path else metadata_json_path
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(metadata_data, f, indent=2, ensure_ascii=False)
    
    print(f"✓ Updated {updated_count} placeholders with LLM contexts")
    print(f"✓ Saved to: {output_file}")
    
    return metadata_data


def generate_and_update_contexts(metadata_path: str, docx_path: str):
    
    print("Generating LLM contexts for placeholders...")
    contexts = generate_placeholder_contexts(metadata_path, docx_path)

    print(f"\nGenerated {len(contexts)} placeholder contexts")
    print("\nFirst context example:")
    if contexts:
        print(json.dumps(contexts[0], indent=2))
        
        # Update the metadata file with contexts
        print("\nUpdating metadata file...")
        update_metadata_with_contexts(metadata_path, contexts)
    else:
        print("No contexts generated")

## Checking area

# FIRST DOCUMENT UPLOAD API CALL (commented out - use FastAPI endpoint instead):
# result = generate_placeholder_metadata(
#     "/Users/rockramsri/Documents/Smart legal filler/Postmoney_Safe_-_Valuation_Cap_Only_-_FINAL-f2a64add6d21039ab347ee2e7194141a4239e364ffed54bad0fe9cf623bf1691_(4).docx", 
#     output_file="metadata_temp.json", 
#     verbose=True
# )
# metadata_path = "/Users/rockramsri/Documents/Smart legal filler/Main-backend/metadata_temp.json"
# docx_path = "/Users/rockramsri/Documents/Smart legal filler/Postmoney_Safe_-_Valuation_Cap_Only_-_FINAL-f2a64add6d21039ab347ee2e7194141a4239e364ffed54bad0fe9cf623bf1691_(4).docx"
# generate_and_update_contexts(metadata_path, docx_path)


# SECOND DOCUMENT chat API CALL

# ************** QUESTION AND FILLING **************



class QuestionResponse(BaseModel):
    """Response containing a question to ask the user."""
    
    question: str = Field(description="The question to ask the user to fill placeholders")
    reasoning: str = Field(description="Brief explanation of why this question is being asked")


def generate_next_question(metadata_json_path: str, docx_path: str) -> dict:
    """
    Generate the next question to ask the user based on unfilled placeholders.
    
    Args:
        metadata_json_path: Path to the placeholder_metadata.json file
        docx_path: Path to the original .docx file
    
    Returns:
        Dictionary with 'question' and 'reasoning' keys, or None if all filled
    """
    # Load the metadata JSON
    with open(metadata_json_path, 'r', encoding='utf-8') as f:
        metadata_data = json.load(f)
    
    # Get all unfilled placeholders
    unfilled_placeholders = [
        p for p in metadata_data['placeholders'] 
        if not p.get('is_filled', False)
    ]
    
    # If all placeholders are filled, return a completion message
    if not unfilled_placeholders:
        return {
            'question': None,
            'reasoning': '🎉🎊 Amazing! All placeholders have been filled successfully! Your document is ready! 🚀✨',
            'status': 'complete'
        }
    
    # Load the document to get full text context
    doc = Document(docx_path)
    
    # Extract full document text for context
    full_document_text = []
    for para in doc.paragraphs:
        text = ''.join([run.text for run in para.runs])
        if text.strip():
            full_document_text.append(text)
    
    document_text_sample = '\n\n'.join(full_document_text[:30])  # First 30 paragraphs for context
    
    # Prepare unfilled placeholders info
    unfilled_info = [
        {
            'unique_id': p['unique_id'],
            'placeholder': p['match'],
            'llm_context': p.get('llm_context'),
            'sentence_with_match': p.get('sentence_with_match'),
            'paragraph_context_before': p.get('paragraph_context_before'),
            'paragraph_context_after': p.get('paragraph_context_after'),
        }
        for p in unfilled_placeholders
    ]
    
    unfilled_json = json.dumps(unfilled_info, indent=2)[:6000]  # Limit size
    
    prompt = f"""You are a friendly, enthusiastic assistant helping to fill out a legal SAFE document! 🎉 Make this process enjoyable and conversational.

    DOCUMENT CONTEXT (first 30 paragraphs):
    {document_text_sample}

    UNFILLED PLACEHOLDERS TO ASK ABOUT:
    {unfilled_json}

    Generate ONE natural, engaging question to ask the user. Vary your phrasing each time - be creative and fun! 😊

    IMPORTANT GUIDELINES:
    1. **Vary Your Questions**: Don't repeat the same phrasing! Mix it up:
       - Sometimes be casual: "Hey! What's the company name?"
       - Sometimes be formal: "Could you provide the company's legal name?"
       - Sometimes combine related info: "What's the company name and which state is it incorporated in?"
    
    2. **Combine Related Fields**: It's OKAY to ask multiple related things in one question if they're connected:
       - "What's the company name and state of incorporation?" ✅
       - "Who's the investor and what's the investment amount?" ✅
       - "What's the date for this SAFE and the purchase amount?" ✅
    
    3. **Make it Fun & Interactive**: 
       - Use friendly emojis naturally (🎯 ✨ 🚀 💼 📝 🎉 😊) but don't overdo it
       - Be enthusiastic and encouraging
       - Use varied greetings: "Let's fill this out!", "Great progress!", "Almost there!"
       - Celebrate when appropriate: "Fantastic! Now let's get..."
    
    4. **Smart Prioritization**:
       - Prioritize the most important/recurring placeholders (like Company Name, Investor Name)
       - Ask about fields that will unlock multiple related fields
       - Group related information naturally
    
    5. **Important Note**: Company names CAN be the same across multiple fields - this is normal! Don't be overly cautious if multiple placeholders ask for company information.
    
    6. **Be Conversational**: Talk like you're helping a friend, not filling out a form!

    EXAMPLES OF GOOD QUESTIONS (varied styles):
    - "Hey! What's the name of the company we're working with? 🏢"
    - "Great progress so far! Could you tell me the investor's name and the investment amount? 💼"
    - "Let's get the basics sorted - what's the company name and which state is it incorporated in? 📝"
    - "Fantastic! Now we need the date for this SAFE agreement. What date should we use? 📅"
    - "Almost there! ✨ Who's the investor in this deal?"
    - "Wonderful! What's the valuation cap and purchase amount? 💰"

    Return a structured response with:
    - question: The engaging question with natural emojis (1-3 max, used tastefully)
    - reasoning: Brief explanation of why you're asking this (1-2 sentences, keep it friendly!)
    """
    
    # Use structured output
    structured_llm = llm.with_structured_output(QuestionResponse)
    
    try:
        response = structured_llm.invoke([{"role": "user", "content": prompt}])
        return {
            'question': response.question,
            'reasoning': response.reasoning,
            'status': 'success',
            'unfilled_count': len(unfilled_placeholders)
        }
    except Exception as e:
        print(f"Error generating question: {e}")
        # Fallback: try without structured output
        response = llm.invoke([{"role": "user", "content": prompt}])
        print("Raw response:")
        print(response.content)
        return {
            'question': "Hey! 😊 Could you tell me what the company name is for this document?",
            'reasoning': 'Error generating question, using fallback - but still excited to help! ✨',
            'status': 'fallback',
            'unfilled_count': len(unfilled_placeholders)
        }





class PlaceholderFill(BaseModel):
    """Information about a placeholder to fill."""
    
    placeholder_id: str = Field(description="The unique_id of the placeholder to fill")
    value: str = Field(description="The value to fill in for this placeholder")
    confidence: str = Field(description="High, Medium, or Low confidence level")
    reasoning: str = Field(description="Brief explanation of why this value fits")


class PlaceholderFillsList(BaseModel):
    """List of placeholder fills from user response."""
    
    fills: List[PlaceholderFill] = Field(
        description="List of placeholders that can be filled from the user's response"
    )


def parse_user_response_and_fill(user_response: str, metadata_json_path: str, docx_path: str) -> dict:
    """
    Parse user response and fill matching placeholders.
    
    Args:
        user_response: The user's response text
        metadata_json_path: Path to the placeholder_metadata.json file
        docx_path: Path to the original .docx file
    
    Returns:
        Dictionary with filling results and updated metadata
    """
    # Load the metadata JSON
    with open(metadata_json_path, 'r', encoding='utf-8') as f:
        metadata_data = json.load(f)
    
    # Get all unfilled placeholders
    unfilled_placeholders = [
        p for p in metadata_data['placeholders'] 
        if not p.get('is_filled', False)
    ]
    
    if not unfilled_placeholders:
        return {
            'status': 'complete',
            'message': 'All placeholders have already been filled!',
            'fills': []
        }
    
    # Load the document to get context
    doc = Document(docx_path)
    
    # Extract full document text
    full_document_text = []
    for para in doc.paragraphs:
        text = ''.join([run.text for run in para.runs])
        if text.strip():
            full_document_text.append(text)
    
    document_text_sample = '\n\n'.join(full_document_text[:30])  # First 30 paragraphs
    
    # Prepare unfilled placeholders info
    unfilled_info = [
        {
            'unique_id': p['unique_id'],
            'placeholder': p['match'],
            'llm_context': p.get('llm_context'),
            'sentence_with_match': p.get('sentence_with_match'),
            'surrounding_text': p.get('surrounding_text'),
        }
        for p in unfilled_placeholders
    ]
    
    unfilled_json = json.dumps(unfilled_info, indent=2)[:6000]  # Limit size
    
    prompt = f"""You are a helpful assistant parsing user responses to fill placeholders in a legal SAFE document.

    DOCUMENT CONTEXT (paragraphs):
    {document_text_sample}

    UNFILLED PLACEHOLDERS:
    {unfilled_json}

    USER RESPONSE:
    "{user_response}"

    Analyze the user's response and determine which placeholders can be filled with the information provided.

    CRITICAL MATCHING RULES:
    1. **Context-Based Matching**: Read the "llm_context" field for each placeholder to understand WHAT it expects. Different placeholders may have the same text (like "[_____________]") but expect different values based on context.

    2. **Smart Multi-Fill**: If user provides company name (e.g., "TechStart Inc."), analyze ALL placeholders to find which ones semantically match based on their llm_context:
    - Fill ALL instances where llm_context indicates "company name" 
    - Fill ALL instances where llm_context indicates "COMPANY" 
    - Fill any other placeholders whose llm_context matches the semantic meaning
    IMPORTANT: Each matching placeholder has a unique placeholder_id. You MUST return a separate entry for EACH matching placeholder.

    3. **Don't Fill Mismatched Context**: If a placeholder text looks similar but the llm_context indicates it expects something different, DO NOT fill it. For example:
    - If [_____________] has llm_context about "Purchase Amount" (money), don't fill it with company name
    - If [_____________] has llm_context about "company name", then fill it with company name
    
    4. **Pattern Recognition**: For ANY piece of information provided, check ALL unfilled placeholders by analyzing their llm_context to determine if they semantically match.

    5. **Confidence**: Only fill if you're reasonably certain the information matches the llm_context expectations

    6. **Extract Exact Values**: Provide specific, exact values extracted from the user's response

    7. **Context is King**: The llm_context field is the authoritative source for what each placeholder expects, not just the placeholder text.

    SPELLING, GRAMMAR, AND FORMATTING RULES:
    8. **Spell-Checking**: Automatically correct any spelling errors in user input before filling placeholders. Examples:
    - "10th octaber" → "October 10" or "10th October" (proper date format)
    - "TechStart Incorperated" → "TechStart Incorporated"
    - "John Doe Compnay" → "John Doe Company"
    - "california" → "California" (proper capitalization for state names)

    9. **Grammar Correction**: Fix grammatical errors and ensure proper formatting:
    - Capitalize proper nouns (company names, person names, state names)
    - Fix common typos and misspellings
    - Normalize date formats (e.g., "octaber 10" → "October 10" or proper date format like "10/10/2024")
    - Correct common word errors ("there" vs "their", "its" vs "it's")

    10. **Value Normalization**: Ensure values are in proper format according to their type:
    - **Dates**: Convert to standard format (MM/DD/YYYY, or full format like "October 10, 2024", or "2024-10-10")
    - **Currency**: Format as numeric value (e.g., "100000" or "100,000" - remove currency symbols if needed)
    - **State Names**: Full state names with proper capitalization (e.g., "California", "Delaware", not "california" or "CALIFORNIA")
    - **Company Names**: Proper capitalization and legal entity suffixes (Inc., LLC, Corp., etc.)
    - **Person Names**: Proper capitalization of first and last names

    11. **Quality Assurance**: Before returning a value, verify:
    - Spelling is correct
    - Grammar is proper
    - Format matches expected type (date, currency, text)
    - Capitalization is appropriate
    - No typos or transposed letters

    EXAMPLE: If user says "company name is TechStart Inc. and date is 10th octaber 2024" and there are:
    - [Company Name] with llm_context: "company name" → FILL with "TechStart Inc." (corrected if needed)
    - [Date of Safe] with llm_context: "date" → FILL with "October 10, 2024" or "10/10/2024" (corrected spelling and proper date format)
    - [_____________] with llm_context: "company name" → FILL with corrected company name
    - [_____________] with llm_context: "Purchase Amount" → DON'T FILL (different context)

    For EACH placeholder that can be filled:
    - placeholder_id: The unique_id from the metadata
    - value: The CORRECTED, SPELL-CHECKED, and PROPERLY FORMATTED value to fill in (NOT the raw user input)
    - confidence: High, Medium, or Low
    - reasoning: Why this value matches this placeholder based on llm_context, and any corrections made (e.g., "Corrected spelling of 'octaber' to 'October' and formatted as proper date")

    Return a structured JSON with the list of fills. Remember: 
    - Match based on llm_context, not placeholder text!
    - Always spell-check, grammar-check, and format values properly before filling!
    """
    
    # Use structured output
    structured_llm = llm.with_structured_output(PlaceholderFillsList)
    
    try:
        response = structured_llm.invoke([{"role": "user", "content": prompt}])
        
        # Process the fills and update metadata
        fills_applied = []
        
        for fill in response.fills:
            # Find the placeholder in metadata
            placeholder_found = False
            for p in metadata_data['placeholders']:
                if p['unique_id'] == fill.placeholder_id:
                    # Update the placeholder
                    p['value'] = fill.value
                    p['is_filled'] = True
                    p['fill_confidence'] = fill.confidence
                    p['fill_reasoning'] = fill.reasoning
                    p['filled_at'] = str(__import__('datetime').datetime.now())
                    placeholder_found = True
                    fills_applied.append({
                        'placeholder_id': fill.placeholder_id,
                        'match': p['match'],
                        'value': fill.value,
                        'confidence': fill.confidence,
                        'reasoning': fill.reasoning
                    })
                    break
            
            if not placeholder_found:
                print(f"Warning: Placeholder ID {fill.placeholder_id} not found in metadata")
        
        # Save updated metadata
        with open(metadata_json_path, 'w', encoding='utf-8') as f:
            json.dump(metadata_data, f, indent=2, ensure_ascii=False)
        
        return {
            'status': 'success',
            'fills_applied': fills_applied,
            'total_fills': len(fills_applied),
            'remaining_unfilled': len([p for p in metadata_data['placeholders'] if not p.get('is_filled', False)])
        }
        
    except Exception as e:
        print(f"Error parsing response: {e}")
        # Fallback: try without structured output
        response = llm.invoke([{"role": "user", "content": prompt}])
        print("Raw response:")
        print(response.content)
        return {
            'status': 'error',
            'message': f'Error parsing response: {str(e)}',
            'fills_applied': []
        }



def fill_and_ask(metadata_path: str, docx_path: str,user_input: str)->dict:
    
    fill_result = parse_user_response_and_fill(user_input, metadata_path, docx_path)
    
    if fill_result['status'] == 'success':
        print(f"\n✓ Filled {fill_result['total_fills']} placeholder(s)")
        for fill in fill_result['fills_applied']:
            print(f"  - {fill['match']}: {fill['value']} ({fill['confidence']} confidence)")
        print(f"Remaining: {fill_result['remaining_unfilled']} placeholders")
    else:
        print(f"Error: {fill_result['message']}")
    
    q_result = generate_next_question(metadata_path, docx_path)
    
    if q_result['status'] == 'complete':
        print("✅ All placeholders filled!")
        return {
            'status': 'complete',
            'message': 'All placeholders filled!',
            'fills': fill_result['fills_applied'],
            'total_fills': fill_result['total_fills'],
            'remaining_unfilled': fill_result['remaining_unfilled'],
            'question': q_result['question'],
            'reasoning': q_result['reasoning']
        }
    else:
        return {
            'status': 'incomplete',
            'message': 'Incomplete placeholders filled',
            'fills': fill_result['fills_applied'],
            'total_fills': fill_result['total_fills'],
            'remaining_unfilled': fill_result['remaining_unfilled'],
            'question': q_result['question'],
            'reasoning': q_result['reasoning']
        }


# Example usage (commented out):
# result = parse_user_response_and_fill(
#     "The company name is TechStart Inc. and the investor is John Smith.",
#     "/Users/rockramsri/Documents/Smart legal filler/placeholder_metadata.json",
#     "/Users/rockramsri/Documents/Smart legal filler/Postmoney_Safe_-_Valuation_Cap_Only_-_FINAL-f2a64add6d21039ab347ee2e7194141a4239e364ffed54bad0fe9cf623bf1691_(4).docx"
# )
# print(result)


# Example usage (commented out):
# result = generate_next_question(
#     "/Users/rockramsri/Documents/Smart legal filler/placeholder_metadata.json",
#     "/Users/rockramsri/Documents/Smart legal filler/Postmoney_Safe_-_Valuation_Cap_Only_-_FINAL-f2a64add6d21039ab347ee2e7194141a4239e364ffed54bad0fe9cf623bf1691_(4).docx"
# )
# print(result)


# FILL AND ASK API CALL (commented out - use FastAPI endpoints instead):
# result = fill_and_ask(
#     "/Users/rockramsri/Documents/Smart legal filler/Main-backend/metadata_temp.json",
#     "/Users/rockramsri/Documents/Smart legal filler/Postmoney_Safe_-_Valuation_Cap_Only_-_FINAL-f2a64add6d21039ab347ee2e7194141a4239e364ffed54bad0fe9cf623bf1691_(4).docx",
#     "The company name is TechStart Inc. and the investor is John Smith."
# )
# print(result)


### Current document downaload API CALL

def fill_document_with_values(metadata_json_path: str, input_docx_path: str, output_docx_path: str) -> dict:
    """
    Fill a Word document with placeholder values from metadata JSON.
    This is the tested version from filter.ipynb.
    
    Args:
        metadata_json_path: Path to placeholder_metadata.json file
        input_docx_path: Path to input .docx file
        output_docx_path: Path to save the filled document
    
    Returns:
        Dictionary with fill statistics and results
    """
    # Load metadata
    with open(metadata_json_path, 'r', encoding='utf-8') as f:
        metadata_data = json.load(f)
    
    # Load document
    doc = Document(input_docx_path)
    
    # Get only filled placeholders
    filled_placeholders = [
        p for p in metadata_data['placeholders'] 
        if p.get('is_filled', False) and p.get('value') is not None
    ]
    
    if not filled_placeholders:
        return {
            'status': 'no_fills',
            'message': 'No filled placeholders found in metadata',
            'total_filled': 0,
            'fills_applied': []
        }
    
    # Group filled placeholders by paragraph index
    fills_by_paragraph = {}
    for p in filled_placeholders:
        para_idx = p.get('paragraph_index')
        if para_idx is not None:
            if para_idx not in fills_by_paragraph:
                fills_by_paragraph[para_idx] = []
            fills_by_paragraph[para_idx].append(p)
    
    fills_applied = []
    total_replaced = 0
    
    # Process each paragraph that has fills
    for para_idx in fills_by_paragraph:
        if para_idx >= len(doc.paragraphs):
            print(f"Warning: Paragraph index {para_idx} out of range")
            continue
        
        paragraph = doc.paragraphs[para_idx]
        
        # Get all fills for this paragraph
        paragraph_fills = fills_by_paragraph[para_idx]
        
        # Sort by position_in_paragraph to replace from right to left (prevent index shifting)
        paragraph_fills.sort(key=lambda x: x.get('position_in_paragraph', 0), reverse=True)
        
        # For each fill in this paragraph
        for fill in paragraph_fills:
            placeholder_text = fill['match']
            replacement_value = fill['value']
            
            # Skip if already replaced (shouldn't happen with reverse sort, but safety check)
            if paragraph_text := ''.join([run.text for run in paragraph.runs]):
                if placeholder_text not in paragraph_text:
                    continue
            
            # Save original run properties
            original_runs = paragraph.runs
            if not original_runs:
                continue
            
            # Get the first run's properties as base styling
            first_run_props = {
                'bold': original_runs[0].bold,
                'italic': original_runs[0].italic,
                'underline': original_runs[0].underline,
                'font_name': original_runs[0].font.name if original_runs[0].font.name else None,
                'font_size': original_runs[0].font.size,
                'font_color': original_runs[0].font.color.rgb if original_runs[0].font.color.rgb else None
            }
            
            # Get full text to replace
            full_text = ''.join([run.text for run in original_runs])
            
            # Replace the placeholder
            if placeholder_text in full_text:
                new_text = full_text.replace(placeholder_text, replacement_value, 1)
                
                # Clear all runs
                for run in paragraph.runs:
                    run.text = ""
                
                # Add the new text with the first run's properties
                if paragraph.runs:
                    paragraph.runs[0].text = new_text
                    # Apply original styling
                    paragraph.runs[0].bold = first_run_props['bold']
                    paragraph.runs[0].italic = first_run_props['italic']
                    paragraph.runs[0].underline = first_run_props['underline']
                    if first_run_props['font_name']:
                        paragraph.runs[0].font.name = first_run_props['font_name']
                    if first_run_props['font_size']:
                        paragraph.runs[0].font.size = first_run_props['font_size']
                    if first_run_props['font_color']:
                        paragraph.runs[0].font.color.rgb = first_run_props['font_color']
                
                fills_applied.append({
                    'placeholder_id': fill['unique_id'],
                    'placeholder': placeholder_text,
                    'value': replacement_value,
                    'paragraph': para_idx
                })
                total_replaced += 1
    
    # Also process tables if needed (for future enhancement)
    table_fills = [p for p in filled_placeholders if p.get('match_type') == 'table']
    if table_fills:
        print(f"Warning: {len(table_fills)} table fills found but not yet implemented")
    
    # Save the filled document
    doc.save(output_docx_path)
    
    return {
        'status': 'success',
        'total_filled': total_replaced,
        'fills_applied': fills_applied,
        'output_path': output_docx_path
    }


# Document ID management
documents_store: Dict[str, Dict[str, str]] = {}
STORAGE_DIR = os.path.join(os.path.dirname(__file__), "document_storage")

# Ensure storage directory exists
os.makedirs(STORAGE_DIR, exist_ok=True)


def create_document_id() -> str:
    """Generate a unique document ID"""
    return str(uuid.uuid4())


def store_document(doc_id: str, original_docx_path: str, metadata_path: str):
    """Store document paths by ID"""
    documents_store[doc_id] = {
        'original_docx_path': original_docx_path,
        'metadata_path': metadata_path,
        'created_at': datetime.now().isoformat()
    }


def get_document_paths(doc_id: str) -> Dict[str, str]:
    """Get document paths by ID"""
    if doc_id not in documents_store:
        raise HTTPException(status_code=404, detail="Document not found")
    return documents_store[doc_id]


### FastAPI Application
app = FastAPI(title="Smart Legal Filler API")

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # In production, specify your frontend URL
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


class ChatRequest(BaseModel):
    """Request model for chat endpoint"""
    user_input: str = Field(description="User's input text to fill placeholders")


@app.post("/upload-document")
async def upload_document(file: UploadFile = File(...)):
    """
    Upload a document and generate metadata.
    Returns document_id for subsequent API calls.
    """
    # Validate file type
    if not file.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="Only .docx files are supported")
    
    # Generate document ID
    doc_id = create_document_id()
    
    # Save uploaded file
    original_docx_path = os.path.join(STORAGE_DIR, f"{doc_id}_original.docx")
    with open(original_docx_path, 'wb') as f:
        content = await file.read()
        f.write(content)
    
    # Generate metadata
    metadata_path = os.path.join(STORAGE_DIR, f"{doc_id}_metadata.json")
    try:
        result = generate_placeholder_metadata(
            original_docx_path,
            output_file=metadata_path,
            verbose=False
        )
        
        # Generate LLM contexts
        generate_and_update_contexts(metadata_path, original_docx_path)
        
        # Store document info
        store_document(doc_id, original_docx_path, metadata_path)
        
        return {
            'status': 'success',
            'document_id': doc_id,
            'message': 'Document uploaded and metadata generated successfully',
            'summary': {
                'total_placeholders': result['summary']['total_placeholders_found'],
                'unique_placeholders': result['summary']['unique_placeholder_count']
            }
        }
    except Exception as e:
        # Clean up on error
        if os.path.exists(original_docx_path):
            os.remove(original_docx_path)
        if os.path.exists(metadata_path):
            os.remove(metadata_path)
        raise HTTPException(status_code=500, detail=f"Error processing document: {str(e)}")


@app.post("/chat/{document_id}")
async def chat_with_document(document_id: str, request: ChatRequest):
    """
    Process user input to fill placeholders and get next question.
    Uses fill_and_ask() function.
    
    Request body: {"user_input": "The company name is TechStart Inc."}
    """
    # Get document paths
    doc_info = get_document_paths(document_id)
    metadata_path = doc_info['metadata_path']
    docx_path = doc_info['original_docx_path']
    
    try:
        # Use fill_and_ask function
        result = fill_and_ask(metadata_path, docx_path, request.user_input)
        
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing chat: {str(e)}")


@app.get("/placeholders/{document_id}")
async def get_placeholders_status(document_id: str):
    """
    Get placeholder status and context information for UI display.
    Returns list of all placeholders with their fill status, LLM context, and summary statistics.
    """
    # Get document paths
    doc_info = get_document_paths(document_id)
    metadata_path = doc_info['metadata_path']
    
    try:
        # Load metadata
        with open(metadata_path, 'r', encoding='utf-8') as f:
            metadata_data = json.load(f)
        
        placeholders = metadata_data.get('placeholders', [])
        
        # Calculate statistics
        total_placeholders = len(placeholders)
        filled_count = sum(1 for p in placeholders if p.get('is_filled', False))
        unfilled_count = total_placeholders - filled_count
        
        # Prepare placeholder list with essential info
        placeholder_list = []
        for p in placeholders:
            llm_context = p.get('llm_context', '')
            # Extract a short summary from LLM context (first 150 chars)
            context_snippet = llm_context[:150] + '...' if len(llm_context) > 150 else llm_context
            
            placeholder_info = {
                'unique_id': p.get('unique_id'),
                'match': p.get('match'),
                'match_type': p.get('match_type'),
                'is_filled': p.get('is_filled', False),
                'value': p.get('value') if p.get('is_filled') else None,
                'llm_context': llm_context,
                'context_snippet': context_snippet if llm_context else 'Context not available',
                'sentence_with_match': p.get('sentence_with_match', '')[:100] if p.get('sentence_with_match') else None,
                'paragraph_index': p.get('paragraph_index'),
                'estimated_page_number': p.get('estimated_page_number'),
                'fill_confidence': p.get('fill_confidence') if p.get('is_filled') else None,
            }
            placeholder_list.append(placeholder_info)
        
        return {
            'status': 'success',
            'summary': {
                'total_placeholders': total_placeholders,
                'filled_count': filled_count,
                'unfilled_count': unfilled_count,
                'completion_percentage': round((filled_count / total_placeholders * 100) if total_placeholders > 0 else 0, 2)
            },
            'placeholders': placeholder_list
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error loading placeholders: {str(e)}")


@app.get("/download/{document_id}")
async def download_document(document_id: str):
    """
    Download the current filled document.
    Returns the updated .docx file with all filled values.
    """
    # Get document paths
    doc_info = get_document_paths(document_id)
    metadata_path = doc_info['metadata_path']
    original_docx_path = doc_info['original_docx_path']
    
    try:
        # Generate filled document
        filled_docx_path = os.path.join(STORAGE_DIR, f"{document_id}_filled.docx")
        fill_result = fill_document_with_values(metadata_path, original_docx_path, filled_docx_path)
        
        # Check if fill was successful
        if fill_result.get('status') == 'no_fills':
            # Still return the original document if no fills
            return FileResponse(
                original_docx_path,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                filename=f"document_{document_id}.docx"
            )
        
        # Return filled document
        return FileResponse(
            filled_docx_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=f"filled_document_{document_id}.docx"
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating document: {str(e)}")